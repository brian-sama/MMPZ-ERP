from pathlib import Path
from shutil import copy2

from docx import Document


SOURCE_DOC = Path(r"C:\Users\brian\Downloads\MMPZ_ERP_Intranet_User_Guide_Draft.docx")
OUTPUT_DOC = Path(r"c:\Users\brian\3D Objects\Personal Projects\MMPZ ERP\output\doc\MMPZ_ERP_Intranet_User_Guide_System_Aligned.docx")


def clear_paragraph(paragraph):
    element = paragraph._element
    for child in list(element):
        if child.tag.endswith("}pPr"):
            continue
        element.remove(child)


def copy_run_format(source_run, target_run):
    if source_run is None:
        return

    target_run.bold = source_run.bold
    target_run.italic = source_run.italic
    target_run.underline = source_run.underline

    if source_run.font.size is not None:
        target_run.font.size = source_run.font.size
    if source_run.font.name:
        target_run.font.name = source_run.font.name
    if source_run.font.color and source_run.font.color.rgb:
        target_run.font.color.rgb = source_run.font.color.rgb


def set_paragraph_segments(paragraph, segments):
    template_runs = list(paragraph.runs)
    clear_paragraph(paragraph)

    if not segments:
        return

    for index, text in enumerate(segments):
        run = paragraph.add_run(text)
        source_run = None
        if template_runs:
            source_run = template_runs[min(index, len(template_runs) - 1)]
        copy_run_format(source_run, run)


def set_cell_text(cell, text):
    if not cell.paragraphs:
        cell.text = text
        return

    set_paragraph_segments(cell.paragraphs[0], [text])
    for extra_paragraph in cell.paragraphs[1:]:
        clear_paragraph(extra_paragraph)


PARAGRAPH_SEGMENTS = {
    0: ["Million Memory Project Zimbabwe"],
    1: ["ERP and Intranet User Guide"],
    2: ["Plain-language manual for leadership, finance, operations, program teams, facilitators, interns, and administrators"],
    4: ["This guide uses simple language so it can be used for training, onboarding, and day-to-day reference."],
    7: ["This guide explains the current MMPZ ERP and intranet in simple, practical language. It tells you what the system is, what each module does, who uses it, and how people in each role work inside the system."],
    8: ["The ERP brings leadership, program delivery, finance, administration, logistics, governance, reports, and field reporting into one system. Instead of splitting work across disconnected spreadsheets and chat threads, the system keeps decisions, records, and operational data in one place."],
    9: ["The intranet is the private internal part of the system. It gives staff, interns, and facilitators a shared space for announcements, contacts, documents, and calendar information."],
    13: ["Different people need different screens and different access. In this system, role codes control which pages a person can open and which tasks they can perform. A user should only see what they need for their work."],
    16: ["Dashboard: ", "The dashboard is the first screen for leadership and oversight users. It shows key numbers, pending approvals, program signals, and shortcuts to the work that needs attention."],
    17: ["Analytics: ", "This area gives a deeper view of trends and performance. It helps leadership and M&E users compare progress, spot patterns, and see where follow-up is needed."],
    18: ["Programs: ", "This is where strategic work is organized into programs, projects, and activities. It helps teams plan delivery, track timelines, and understand how field work is grouped."],
    19: ["Facilitators: ", "This area helps program staff manage development facilitators, assignments, coverage, and operational coordination for work happening in the field."],
    20: ["Monitoring and Evaluation: ", "This is where indicators, implementation progress, and field results are reviewed. It supports evidence-based reporting instead of guesswork."],
    21: ["Finance, Administration and Logistics: ", "This is where grants, budgets, budget lines, commitments, and procurement requisitions are managed. It supports disciplined spending, documented approvals, and operational readiness."],
    22: ["Governance Queue: ", "This is the approval center. Reviewers open pending items here, check the control notes and item details, and approve or reject transactions with an audit trail."],
    23: ["Reports: ", "This area generates management and donor-facing outputs. Users can choose a report template, set the period, and download the result in PDF or spreadsheet format."],
    24: ["Settings: ", "This is where core operational rules are maintained, including the major finance threshold that controls when a request escalates for higher approval."],
    25: ["User Management: ", "This area is used for account setup, role assignment, and access administration so each person sees the right parts of the system."],
    26: ["My Portal: ", "This is the field-facing workspace for facilitators and some interns. It supports assignments, daily attendance, narrative submissions, and activity reporting."],
    27: ["Intranet: ", "This is the shared internal information space for announcements, contacts, documents, and calendar information across the organization."],
    29: ["5.1 System Administrator"],
    30: ["In simple terms: this person keeps the system available, secure, and under control."],
    31: ["• ", "Creates user accounts and assigns the correct level of access."],
    32: ["• ", "Supports login, password reset, and account access problems."],
    33: ["• ", "Maintains technical settings, backups, and environment checks."],
    34: ["• ", "Reviews logs and unusual activity so problems can be traced quickly."],
    35: ["• ", "Protects integrations and configuration that keep the system running."],
    36: ["5.2 Director"],
    37: ["In simple terms: this person sees the organizational picture and the final control points."],
    38: ["• ", "Looks at dashboards, analytics, and reports to understand overall performance."],
    39: ["• ", "Approves major requests, policy thresholds, and high-value transactions."],
    40: ["• ", "Checks progress across programs, finance, field delivery, and pending decisions."],
    41: ["• ", "Acts when risks, delays, or overspending signals appear."],
    42: ["5.3 Finance Admin Officer"],
    43: ["In simple terms: this person protects financial accuracy and spending discipline."],
    44: ["• ", "Reviews grants, budgets, budget lines, and utilization levels."],
    45: ["• ", "Checks requisitions before money is committed or ordering starts."],
    46: ["• ", "Monitors commitments, pending approvals, and available balance."],
    47: ["• ", "Produces finance reports and supports donor accountability."],
    48: ["5.4 Admin Assistant"],
    49: ["In simple terms: this person keeps administrative work orderly and moving."],
    50: ["• ", "Coordinates records, settings, and routine administration."],
    51: ["• ", "Supports user setup and follow-up on internal workflows."],
    52: ["• ", "Prepares documentation needed for approvals and audits."],
    53: ["• ", "Uses intranet and reports to keep staff informed and aligned."],
    54: ["5.5 Logistics Assistant"],
    55: ["In simple terms: this person helps move approved purchases into real delivery."],
    56: ["• ", "Reviews approved requisitions and the delivery requirements behind them."],
    57: ["• ", "Tracks ordering, movement, and receiving of goods."],
    58: ["• ", "Checks that logistics actions follow approved, budget-backed requests."],
    59: ["• ", "Flags delays, missing documents, or receiving issues before they grow."],
    60: ["5.6 Community Development Officer"],
    61: ["In simple terms: this person plans and tracks community implementation work."],
    62: ["• ", "Builds or updates programs, projects, and field activities."],
    63: ["• ", "Coordinates facilitators and monitors progress in delivery areas."],
    64: ["• ", "Initiates requisitions when field work needs approved resources."],
    65: ["• ", "Uses reports and indicators to see whether interventions are moving as planned."],
    66: ["5.7 Psychosocial Support Officer"],
    67: ["In simple terms: this person manages service delivery and case-support activities."],
    68: ["• ", "Plans support activities and records implementation progress."],
    69: ["• ", "Coordinates facilitators and monitors attendance where relevant."],
    70: ["• ", "Raises operational requests for approved service delivery needs."],
    71: ["• ", "Watches program indicators and reports for quality and follow-up."],
    72: ["5.8 M&E Intern Acting Officer"],
    73: ["In simple terms: this person turns field information into usable evidence."],
    74: ["• ", "Reviews indicators, counts, and reporting submissions."],
    75: ["• ", "Uses analytics and reports to compare progress against targets."],
    76: ["• ", "Checks data quality and follows up on missing or unusual entries."],
    77: ["• ", "Helps management understand trends from field work and approvals."],
    78: ["5.9 Development Facilitator"],
    79: ["In simple terms: this person carries the work into the field."],
    80: ["• ", "Checks assignments in My Portal."],
    81: ["• ", "Logs daily attendance for field or office work."],
    82: ["• ", "Submits activity counts and weekly narrative updates."],
    83: ["• ", "Keeps implementation notes current so supervisors can review progress."],
    84: ["5.10 Social Services Intern"],
    85: ["In simple terms: this person supports implementation and reporting under supervision."],
    86: ["• ", "Uses My Portal for assigned work and simple reporting."],
    87: ["• ", "Records field activity information using guided forms."],
    88: ["• ", "Reads announcements, documents, and calendar updates on the intranet."],
    89: ["• ", "Escalates issues or data gaps to supervisors instead of guessing."],
    90: ["5.11 Youth Communications Intern"],
    91: ["In simple terms: this person supports youth-facing delivery and communication tasks."],
    92: ["• ", "Uses My Portal for assigned activities and reporting."],
    93: ["• ", "Follows intranet announcements and document guidance."],
    94: ["• ", "Supports communication-related field work and progress updates."],
    95: ["• ", "Shares information through the right channels instead of disconnected records."],
    96: ["5.12 General Staff User"],
    97: ["In simple terms: this person uses the intranet for day-to-day coordination."],
    98: ["• ", "Reads announcements and policy notices."],
    99: ["• ", "Finds documents, contacts, and calendar information."],
    100: ["• ", "Uses the system as a trusted reference point for current organizational information."],
    102: ["This table shows a simple version of the access plan. Full means direct working access, view means read-only or oversight access, limited means a narrower part of the process, and no means the area is normally hidden."],
    104: ["Procurement requisition"],
    105: ["• ", "A requester opens Finance, Administration and Logistics and starts a new requisition."],
    106: ["• ", "The request is linked to the correct budget line and itemized cost lines are added."],
    107: ["• ", "The system calculates the total and assigns a routine, finance-review, or director-review band."],
    108: ["• ", "The requisition enters governance review and waits for the right reviewer."],
    109: ["• ", "Once approved, the amount becomes a commitment and logistics can move toward ordering and delivery."],
    110: ["Governance approval"],
    111: ["• ", "A reviewer opens the governance queue and selects the pending item."],
    112: ["• ", "The reviewer checks the requester, item lines, budget link, and control note."],
    113: ["• ", "The reviewer approves or rejects the item with comments."],
    114: ["• ", "The system records the action in the approval log and updates the requisition status."],
    115: ["Facilitator attendance and reporting"],
    116: ["• ", "The facilitator opens My Portal and selects the assigned work item."],
    117: ["• ", "Attendance is logged for the work date and status."],
    118: ["• ", "Narrative or activity data is submitted from the same portal."],
    119: ["• ", "Supervisors and M&E users later review the entries through reports and indicators."],
    120: ["Intranet notice or document access"],
    121: ["• ", "A user opens the intranet area and reads the latest announcement or searches the directory."],
    122: ["• ", "The user opens the document library or calendar for the needed item."],
    123: ["• ", "Content is shown according to the user's internal access level."],
    124: ["• ", "Staff use the intranet as the formal source for current notices, documents, and organization dates."],
    126: ["The system should feel simple, calm, and professional. A user should not need to guess where to click. Important numbers should be visible, approval steps should be clear, and forms should guide people to complete work correctly the first time."],
    127: ["Example screen ideas are shown below. They reflect the direction of the current ERP and intranet, especially the dashboard, finance workspace, and intranet home area."],
    129: ["Example executive dashboard"],
    131: ["Example finance and logistics workspace"],
    133: ["Example intranet announcements page"],
    137: ["Some information should only be seen by the right people. Finance records, approvals, staff details, field reporting data, donor information, and internal operational documents should not be open to everyone."],
    138: ["• ", "Use strong passwords and do not share accounts."],
    139: ["• ", "Use role-based access so people only see the pages and actions meant for their job."],
    140: ["• ", "Follow maker-checker rules so the requester does not approve their own transaction."],
    141: ["• ", "Keep audit logs so the team can trace decisions, approvals, and changes."],
    142: ["• ", "Link requisitions to approved budget lines before money is committed or goods are ordered."],
    143: ["• ", "Handle staff, donor, and program data according to privacy rules and only export what is truly needed."],
    146: ["This manual can grow with the system. As more features are completed, each role can get its own step-by-step pages, live screenshots, and training notes. The goal is to make the system easy to understand, easy to trust, and easy to use."],
}


TABLE_DATA = {
    0: [
        ["Document purpose", "Explain the current system in simple terms and show how each role uses it."],
        ["Audience", "Leadership, finance, operations, program teams, facilitators, interns, and administrators."],
        ["Version", "System-aligned guide for v1.0.0"],
        ["Prepared for", "Million Memory Project Zimbabwe"],
    ],
    1: [
        ["Main part", "Simple explanation"],
        ["Dashboard", "A home screen for key numbers, alerts, approvals, and shortcuts."],
        ["Programs", "A place to organize programs, projects, and activities."],
        ["Facilitators", "A place to manage facilitator assignments and field coordination."],
        ["Monitoring and Evaluation", "A place to review indicators, field results, and evidence for reporting."],
        ["Finance, Administration and Logistics", "A place for grants, budgets, commitments, requisitions, and operational controls."],
        ["Governance Queue", "A place where approvals are reviewed, actioned, and logged."],
        ["Reports", "A place to generate PDF or spreadsheet exports for management and donors."],
        ["Intranet", "A private staff area for announcements, contacts, documents, and calendar information."],
    ],
    2: [
        ["Role", "What they do", "What they mainly need"],
        ["System Administrator", "Sets up access, resolves account issues, and protects the technical environment.", "User accounts, permissions, settings, logs."],
        ["Director", "Sees the overall picture, reviews major approvals, and checks strategic signals.", "Dashboard, approvals, reports, settings."],
        ["Finance Admin Officer", "Manages grants, budgets, commitments, requisitions, and controls.", "Finance workspace, governance queue, reports."],
        ["Admin Assistant", "Supports administration, records, user setup, and workflow follow-up.", "Settings, user management, governance support, intranet."],
        ["Logistics Assistant", "Follows approved requisitions into ordering, movement, and receiving.", "Finance and logistics workspace, operational follow-up."],
        ["Community Development Officer", "Plans community work, coordinates delivery, and requests resources.", "Programs, facilitators, M&E, requisition access."],
        ["Psychosocial Support Officer", "Manages service delivery activities and related follow-up.", "Programs, facilitators, M&E, requisition access."],
        ["M&E Intern Acting Officer", "Reviews reporting data and helps turn field results into evidence.", "M&E, analytics, reports."],
        ["Development Facilitator", "Implements field work and submits attendance and activity updates.", "My Portal, assignments, attendance, reporting."],
        ["Social Services Intern", "Supports implementation and reporting under supervision.", "My Portal, intranet, guided reporting."],
        ["Youth Communications Intern", "Supports communication-related field work and progress updates.", "My Portal, intranet, communication guidance."],
        ["General Staff User", "Uses the intranet as a reliable source of notices and documents.", "Announcements, directory, documents, calendar."],
    ],
    3: [
        ["Module", "Sys Admin", "Director", "Finance & Ops", "Program Leads", "Field & Interns", "All Users"],
        ["Dashboard", "Full", "Full", "View", "View", "Limited", "No"],
        ["Programs & Facilitators", "Full", "View", "Limited", "Full", "Limited", "No"],
        ["Monitoring & Evaluation", "Full", "View", "Limited", "Full", "Limited", "No"],
        ["Finance & Logistics", "Full", "Approve/View", "Full", "Limited", "No", "No"],
        ["Governance Queue", "Full", "Full", "Review", "No", "No", "No"],
        ["Reports", "Full", "Full", "Full", "View", "Limited", "No"],
        ["Intranet", "Full", "View", "View", "View", "View", "Staff"],
    ],
    4: [
        ["Task", "How it works"],
        ["Logging in", "Go to the system link, enter your username and password, and sign in. After login, the system sends you to the page that matches your role."],
        ["Checking a dashboard", "Open Executive Dashboard or Analytics to review key numbers, alerts, approvals, and performance signals."],
        ["Submitting a requisition", "Open Finance, click New Requisition, choose the budget line, add the justification and item lines, then submit for review."],
        ["Reviewing an approval", "Open Governance Queue, open the item, read the control note and item list, then approve or reject with comments."],
        ["Logging facilitator attendance", "Open My Portal, choose the assignment, select the work date and attendance status, then save the record."],
        ["Generating a report", "Open Reports, choose a report template, set the date range and format, then generate the download."],
        ["Finding a document or announcement", "Open Intranet, read the announcement feed, search the directory, or open the document library and calendar."],
    ],
    5: [
        ["Word", "Meaning"],
        ["ERP", "A system that brings leadership, programs, finance, operations, and reporting into one place."],
        ["Intranet", "A private internal area only for approved users."],
        ["Budget line", "A defined funding line used to control where spending can be charged."],
        ["Commitment", "Money reserved by an approved or ordered requisition before cash is actually paid out."],
        ["Approval band", "The control level a requisition falls into, such as routine review, finance review, or Director review."],
        ["Governance queue", "The approval screen where pending decisions are reviewed and actioned."],
        ["Facilitator portal", "The field-facing page used for assignments, attendance, and reporting."],
        ["Audit trail", "A record of who did what, when they did it, and what decision was made."],
    ],
}


def main():
    if not SOURCE_DOC.exists():
        raise FileNotFoundError(f"Source guide not found: {SOURCE_DOC}")

    OUTPUT_DOC.parent.mkdir(parents=True, exist_ok=True)
    copy2(SOURCE_DOC, OUTPUT_DOC)

    doc = Document(OUTPUT_DOC)

    for paragraph_index, segments in PARAGRAPH_SEGMENTS.items():
        set_paragraph_segments(doc.paragraphs[paragraph_index], segments)

    for table_index, rows in TABLE_DATA.items():
        table = doc.tables[table_index]
        for row_index, row_values in enumerate(rows):
            row = table.rows[row_index]
            for cell_index, value in enumerate(row_values):
                set_cell_text(row.cells[cell_index], value)

    doc.save(OUTPUT_DOC)

    validation = Document(OUTPUT_DOC)
    print(f"Saved: {OUTPUT_DOC}")
    print(f"Paragraphs: {len(validation.paragraphs)}")
    print(f"Tables: {len(validation.tables)}")
    print(f"Title: {validation.paragraphs[1].text}")
    print(f"Section 5 heading: {validation.paragraphs[28].text}")
    print(f"Final heading: {validation.paragraphs[145].text}")


if __name__ == "__main__":
    main()
